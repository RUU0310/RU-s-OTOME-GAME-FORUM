from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 用例说明内容（可根据实际补充/调整）
use_cases = [
    {
        'name': '浏览游戏列表',
        'desc': '用户在游戏一览页面浏览所有已收录的乙女游戏，支持按条件筛选、排序、分页。',
        'actor': '普通用户、游客',
        'pre': '用户已进入"游戏一览"页面（@GameOverViews.vue）。',
        'main': [
            '1. 用户进入"游戏一览"页面。',
            '2. 系统展示所有已收录的乙女游戏列表（卡片形式，含封面、名称、简介、评分等）。',
            '3. 用户可通过筛选条件（如平台、类型、评分等）筛选游戏。',
            '4. 用户可按评分、时间等排序游戏列表。',
            '5. 用户可分页浏览全部游戏。',
            '6. 用户点击某一游戏卡片，进入游戏详情页。'
        ],
        'alt': [
            '3a. 若无符合筛选条件的游戏，系统提示"暂无数据"。'
        ],
        'ex': [
            '2a. 游戏数据加载失败，系统提示"加载失败，请重试"。'
        ],
        'note': '支持游客访问，部分高级筛选需登录。'
    },
    {
        'name': '查看游戏详情',
        'desc': '用户点击游戏卡片，查看该游戏的详细信息，包括介绍、角色、评分、评论等。',
        'actor': '普通用户、游客',
        'pre': '用户已在"游戏一览"页面，且已选中某一游戏。',
        'main': [
            '1. 用户点击游戏卡片。',
            '2. 系统跳转至游戏详情页（@GameDetail.vue），展示该游戏的详细信息：',
            '   - 游戏封面、名称、简介、发售时间、平台、标签等',
            '   - 可攻略角色列表（含图片、姓名、简介）',
            '   - 游戏评分（综合、各维度）',
            '   - 用户评论列表',
            '3. 用户可切换角色卡片，查看角色详情。',
            '4. 用户可查看/参与评论、点赞、收藏等操作（需登录）。'
        ],
        'alt': [
            '2a. 若游戏无可攻略角色，则角色区域显示"暂无角色信息"。',
            '2b. 若无评论，则评论区显示"暂无评论"。'
        ],
        'ex': [
            '2c. 游戏详情加载失败，系统提示"加载失败，请重试"。'
        ],
        'note': '部分操作（评论、点赞、收藏）需登录。'
    },
    # 可继续补充其他用例
]

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return p

def add_paragraph(doc, title, content):
    p = doc.add_paragraph()
    run = p.add_run(f"{title}")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(f"{content}")
    return p

def add_list(doc, title, items):
    p = doc.add_paragraph()
    run = p.add_run(f"{title}")
    run.bold = True
    run.font.size = Pt(11)
    for item in items:
        doc.add_paragraph(item, style='List Number')

def main():
    doc = Document()
    doc.add_heading('游戏一览模块 用例说明', 0)
    doc.add_paragraph('本说明文档详细描述了乙女游戏社区"游戏一览"模块的主要用例，包括用例名称、描述、参与者、前置条件、基本流程、其他流程、异常流程及注释。')
    for uc in use_cases:
        add_heading(doc, f"用例：{uc['name']}", level=1)
        add_paragraph(doc, '用例描述：', uc['desc'])
        add_paragraph(doc, '参与者：', uc['actor'])
        add_paragraph(doc, '前置条件：', uc['pre'])
        add_list(doc, '基本流程：', uc['main'])
        if uc['alt']:
            add_list(doc, '其他流程：', uc['alt'])
        if uc['ex']:
            add_list(doc, '异常流程：', uc['ex'])
        add_paragraph(doc, '注释：', uc['note'])
        doc.add_paragraph('')
    doc.save('UseCase_Description_Games.docx')
    print('文档已生成：UseCase_Description_Games.docx')

if __name__ == '__main__':
    main() 